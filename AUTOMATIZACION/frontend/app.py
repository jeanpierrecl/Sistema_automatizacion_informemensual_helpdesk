from __future__ import annotations

import html
import importlib.util
import json
import copy
import mimetypes
import os
import re
import shutil
import sys
import tempfile
import traceback
import time
import uuid
from email import policy
from email.parser import BytesParser
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, urlparse
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta


APP_DIR = Path(__file__).resolve().parent
WORKSPACE_DIR = APP_DIR.parent
SCRIPT_PATH = APP_DIR / "automatizar_tablas_y_comparativo.py"
TEMPLATE_DIR = APP_DIR / "templates"
STATIC_DIR = APP_DIR / "static"
OUTPUT_DIR = Path(tempfile.gettempdir()) / "automatizacion_reporte_mensual_downloads"
UPLOAD_DIR = WORKSPACE_DIR / ".frontend_uploads"
REPORT_TEMPLATE_PATH = WORKSPACE_DIR.parent / "PLANTILLAS" / "Informe de Consumos - Plantilla.docx"
HELPDESK_TEMPLATE_PATH = WORKSPACE_DIR.parent / "PLANTILLAS" / "XM Soporte_ Informe de Helpdesk- Inspira IT_Mes_Anio - plantilla.docx"
NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
NS = "{%s}" % NS_MAIN
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
MESES_ABREV = {
    1: "Ene",
    2: "Feb",
    3: "Mar",
    4: "Abr",
    5: "May",
    6: "Jun",
    7: "Jul",
    8: "Ago",
    9: "Sep",
    10: "Oct",
    11: "Nov",
    12: "Dic",
}
MESES_NUM = {v: k for k, v in MESES_ABREV.items()}
MESES_LARGO = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}
DOCX_IMAGE_TARGETS = {
    "ambiente": "word/media/image7.png",
    "total": "word/media/image5.png",
}
DOCX_RESOURCE_IMAGE_TARGET = "word/media/image4.png"


def cargar_automatizacion():
    spec = importlib.util.spec_from_file_location("automatizar_tablas_y_comparativo", SCRIPT_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"No se pudo cargar el script: {SCRIPT_PATH}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


AUTOMATIZACION = cargar_automatizacion()


def safe_filename(name: str, fallback: str) -> str:
    clean = "".join(ch for ch in name if ch.isalnum() or ch in " ._-()").strip()
    return clean or fallback


def parse_multipart(handler: BaseHTTPRequestHandler):
    content_type = handler.headers.get("Content-Type", "")
    content_length = int(handler.headers.get("Content-Length", "0"))
    body = handler.rfile.read(content_length)
    message = BytesParser(policy=policy.default).parsebytes(
        f"Content-Type: {content_type}\r\nMIME-Version: 1.0\r\n\r\n".encode("utf-8") + body
    )

    fields: dict[str, str] = {}
    files: dict[str, tuple[str, bytes] | list[tuple[str, bytes]]] = {}
    for part in message.iter_parts():
        if part.get_content_disposition() != "form-data":
            continue
        name = part.get_param("name", header="content-disposition")
        filename = part.get_filename()
        payload = part.get_payload(decode=True) or b""
        if filename:
            item = (safe_filename(filename, f"{name}.xlsx"), payload)
            if name == "resource_images":
                files.setdefault(name, [])
                files[name].append(item)
            else:
                files[name] = item
        else:
            fields[name] = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
    return fields, files


def load_env_file():
    try:
        from dotenv import load_dotenv
    except ImportError:
        return
    for env_path in (WORKSPACE_DIR / ".env", WORKSPACE_DIR.parent / ".env"):
        if env_path.exists():
            load_dotenv(env_path, override=False)


load_env_file()


def save_upload(
    files: dict[str, tuple[str, bytes] | list[tuple[str, bytes]]],
    key: str,
    run_upload_dir: Path,
    required: bool,
    extensions: tuple[str, ...] = (".xlsx", ".xlsm"),
) -> Path | None:
    item = files.get(key)
    if isinstance(item, list):
        item = item[0] if item else None
    if not item or not item[1]:
        if required:
            raise ValueError(f"Falta el archivo requerido: {key}.")
        return None

    filename, content = item
    if not filename.lower().endswith(extensions):
        allowed = ", ".join(extensions)
        raise ValueError(f"El archivo {filename} debe ser {allowed}.")

    path = run_upload_dir / filename
    path.write_bytes(content)
    return path


def save_uploads(
    files: dict[str, tuple[str, bytes] | list[tuple[str, bytes]]],
    key: str,
    run_upload_dir: Path,
    extensions: tuple[str, ...],
) -> list[Path]:
    items = files.get(key)
    if not items:
        return []
    if not isinstance(items, list):
        items = [items]
    paths = []
    for index, (filename, content) in enumerate(items, 1):
        if not content:
            continue
        if not filename.lower().endswith(extensions):
            allowed = ", ".join(extensions)
            raise ValueError(f"El archivo {filename} debe ser {allowed}.")
        path = run_upload_dir / f"{index:02d}_{filename}"
        path.write_bytes(content)
        paths.append(path)
    return paths


def infer_month_from_name(filename: str) -> str | None:
    match = re.search(r"(20\d{2})[-_](0[1-9]|1[0-2])", filename)
    if not match:
        return None
    year = int(match.group(1))
    month = int(match.group(2))
    return f"{MESES_ABREV[month]}-{str(year)[-2:]}"


def resolve_month(fields: dict[str, str], files: dict[str, tuple[str, bytes]]) -> str:
    manual_month = (fields.get("mes_manual") or "").strip()
    inferred = {
        key: infer_month_from_name(item[0])
        for key, item in files.items()
        if key in {"summary", "instances"} and item
    }
    inferred_values = {value for value in inferred.values() if value}
    if len(inferred_values) == 1:
        return inferred_values.pop()
    if len(inferred_values) > 1:
        raise ValueError("Los archivos Summary e Instances tienen meses distintos en el nombre.")
    if manual_month:
        AUTOMATIZACION.parse_etiqueta_mes(manual_month)
        return manual_month
    raise ValueError("No se pudo reconocer el mes desde los nombres de los archivos. Escribe el mes manualmente.")


def month_filename_parts(etiqueta_mes: str) -> tuple[str, str]:
    year, month, _ = AUTOMATIZACION.parse_etiqueta_mes(etiqueta_mes)
    return f"{year}-{month:02d}", f"{MESES_ABREV[month]}-{year}"


def cell_to_month_label(value) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if text.startswith("Total "):
        text = text.replace("Total ", "", 1).strip()
    match = re.fullmatch(r"([A-ZÃÃ‰ÃÃ“ÃšÃ‘a-zÃ¡Ã©Ã­Ã³ÃºÃ±]{3})-(\d{2}|\d{4})", text)
    if not match:
        return None
    abbr = match.group(1).capitalize()
    if abbr not in MESES_NUM:
        return None
    year = int(match.group(2))
    if year < 100:
        year += 2000
    return f"{abbr}-{year}"


def comparativo_first_month(path: Path) -> str | None:
    with zipfile.ZipFile(path, "r") as z:
        shared = workbook_shared_strings(z)
        workbook = ET.fromstring(z.read("xl/workbook.xml"))
        rels = ET.fromstring(z.read("xl/_rels/workbook.xml.rels"))
        relmap = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels}
        ns = {"m": NS_MAIN, "r": NS_R}
        resumen_path = None
        for sheet in workbook.find("m:sheets", ns):
            if sheet.attrib["name"] == "Resumen":
                resumen_path = "xl/" + relmap[sheet.attrib[f"{{{NS_R}}}id"]].lstrip("/")
                break
        if not resumen_path:
            return None
        root = ET.fromstring(z.read(resumen_path))
        rows = {int(row.attrib["r"]): row for row in root.find(NS + "sheetData").findall(NS + "row")}
        row4 = rows.get(4)
        if row4 is None:
            return None
        labels = []
        for cell in row4.findall(NS + "c"):
            label = cell_to_month_label(workbook_cell_value(cell, shared))
            if label:
                year, month, _ = AUTOMATIZACION.parse_etiqueta_mes(label)
                labels.append((year, month, label))
        if not labels:
            return None
        return min(labels)[2]


def workbook_shared_strings(z: zipfile.ZipFile):
    try:
        root = ET.fromstring(z.read("xl/sharedStrings.xml"))
    except KeyError:
        return []
    return ["".join(t.text or "" for t in si.iter(NS + "t")) for si in root.findall(NS + "si")]


def workbook_cell_value(cell, shared):
    if cell is None:
        return None
    cell_type = cell.attrib.get("t")
    value = cell.find(NS + "v")
    inline = cell.find(NS + "is")
    if cell_type == "s" and value is not None:
        return shared[int(value.text)]
    if cell_type == "inlineStr" and inline is not None:
        return "".join(t.text or "" for t in inline.iter(NS + "t"))
    if value is not None:
        return value.text
    return None


def workbook_sheet_path(z: zipfile.ZipFile, sheet_name: str) -> str | None:
    workbook = ET.fromstring(z.read("xl/workbook.xml"))
    rels = ET.fromstring(z.read("xl/_rels/workbook.xml.rels"))
    relmap = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels}
    ns = {"m": NS_MAIN, "r": NS_R}
    for sheet in workbook.find("m:sheets", ns):
        if sheet.attrib["name"] == sheet_name:
            return "xl/" + relmap[sheet.attrib[f"{{{NS_R}}}id"]].lstrip("/")
    return None


def column_index_from_ref(cell_ref: str) -> int:
    letters = re.match(r"[A-Z]+", cell_ref.upper())
    if not letters:
        return 0
    index = 0
    for char in letters.group(0):
        index = index * 26 + ord(char) - ord("A") + 1
    return index


def worksheet_row_maps(z: zipfile.ZipFile, sheet_path: str, max_row: int | None = None):
    root = ET.fromstring(z.read(sheet_path))
    rows = {}
    sheet_data = root.find(NS + "sheetData")
    if sheet_data is None:
        return rows
    for row in sheet_data.findall(NS + "row"):
        row_idx = int(row.attrib.get("r", "0"))
        if max_row is not None and row_idx > max_row:
            continue
        cells = {}
        for cell in row.findall(NS + "c"):
            col_idx = column_index_from_ref(cell.attrib.get("r", ""))
            if col_idx:
                cells[col_idx] = cell
        rows[row_idx] = cells
    return rows


def as_float(value) -> float:
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip()
    if not text or text == "-":
        return 0.0
    text = text.replace("$", "").replace(",", "")
    try:
        return float(text)
    except ValueError:
        return 0.0


def comparativo_has_month(path: Path, etiqueta_mes: str) -> bool:
    with zipfile.ZipFile(path, "r") as z:
        shared = workbook_shared_strings(z)
        workbook = ET.fromstring(z.read("xl/workbook.xml"))
        rels = ET.fromstring(z.read("xl/_rels/workbook.xml.rels"))
        relmap = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels}
        ns = {"m": NS_MAIN, "r": NS_R}
        resumen_path = None
        for sheet in workbook.find("m:sheets", ns):
            if sheet.attrib["name"] == "Resumen":
                resumen_path = "xl/" + relmap[sheet.attrib[f"{{{NS_R}}}id"]].lstrip("/")
                break
        if not resumen_path:
            return False
        root = ET.fromstring(z.read(resumen_path))
        rows = {int(row.attrib["r"]): row for row in root.find(NS + "sheetData").findall(NS + "row")}
        row4 = rows.get(4)
        if row4 is None:
            return False
        return any(
            str(workbook_cell_value(cell, shared)).strip() == f"Total {etiqueta_mes}"
            for cell in row4.findall(NS + "c")
        )


def month_text_parts(etiqueta_mes: str) -> dict[str, str | int]:
    year, month, _ = AUTOMATIZACION.parse_etiqueta_mes(etiqueta_mes)
    return {
        "year": year,
        "month": month,
        "abbr_short": f"{MESES_ABREV[month]}-{str(year)[-2:]}",
        "abbr_long": f"{MESES_ABREV[month]}-{year}",
        "name": MESES_LARGO[month],
        "name_cap": MESES_LARGO[month].capitalize(),
        "title": f"{MESES_ABREV[month]} {year}",
    }


def excel_date_to_datetime(value):
    if isinstance(value, datetime):
        return value
    if isinstance(value, (int, float)):
        return datetime(1899, 12, 30) + timedelta(days=float(value))
    return None


def chart_export_paths(comparativo_path: Path, output_dir: Path, months: list[dict] | None = None) -> dict[str, Path]:
    # Evita depender de Excel COM durante la generacion del Word. Si Excel queda
    # ocupado por un modal o por recalculo de graficos, el informe puede tardar
    # varios minutos; el fallback se arma directo desde la hoja Cuadro.
    return generate_chart_images_fallback(comparativo_path, output_dir, months)


def chart_font(size: int, bold: bool = False):
    from PIL import ImageFont

    names = ("arialbd.ttf", "arial.ttf") if bold else ("arial.ttf", "calibri.ttf")
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def nice_axis_max(value: float) -> float:
    if value <= 0:
        return 1
    step = 5000 if value > 50000 else 2000
    return ((int(value / step) + 1) * step)


def draw_centered(draw, xy, text: str, font, fill="#555555"):
    x, y = xy
    bbox = draw.textbbox((0, 0), text, font=font)
    draw.text((x - (bbox[2] - bbox[0]) / 2, y), text, font=font, fill=fill)


def generate_chart_images_fallback(comparativo_path: Path, output_dir: Path, months: list[dict] | None = None) -> dict[str, Path]:
    try:
        from PIL import Image, ImageDraw
    except ImportError as exc:
        raise ValueError("Falta instalar Pillow. Ejecuta: pip install -r requirements.txt") from exc

    months = months or collect_cuadro_data(comparativo_path)
    if not months:
        raise ValueError("No se pudieron generar los graficos: no hay datos en la hoja Cuadro.")

    ambiente_path = output_dir / "grafico_consumos_por_ambiente.png"
    total_path = output_dir / "grafico_consumos_totales_por_mes.png"
    first = months[0]["date"]
    last = months[-1]["date"]
    last_title = f"{MESES_ABREV[first.month]} {first.year} - {MESES_ABREV[last.month]} {last.year}"

    title_font = chart_font(25, True)
    label_font = chart_font(13)
    small_font = chart_font(11)
    value_font = chart_font(12, True)

    # Stacked environment chart
    img = Image.new("RGB", (1024, 479), "white")
    draw = ImageDraw.Draw(img)
    draw_centered(draw, (512, 24), "Consumos por Ambiente", title_font)
    plot_left, plot_top, plot_right, plot_bottom = 100, 80, 930, 385
    max_total = nice_axis_max(max(m["total"] for m in months))
    colors = {"dev": "#4472C4", "qa": "#ED7D31", "prod": "#A5A5A5"}
    for tick in range(0, int(max_total) + 1, 2000):
        y = plot_bottom - (tick / max_total) * (plot_bottom - plot_top)
        draw.line((plot_left, y, plot_right, y), fill="#D9D9D9", width=1)
        draw.text((18, y - 8), f"${tick:,.2f}", font=small_font, fill="#555555")
    gap = (plot_right - plot_left) / max(len(months), 1)
    bar_w = min(44, gap * 0.5)
    for idx, item in enumerate(months):
        x0 = plot_left + gap * idx + gap / 2 - bar_w / 2
        x1 = x0 + bar_w
        current_bottom = plot_bottom
        for key in ("dev", "qa", "prod"):
            h = (item[key] / max_total) * (plot_bottom - plot_top)
            y0 = current_bottom - h
            draw.rectangle((x0, y0, x1, current_bottom), fill=colors[key])
            if h > 22:
                draw_centered(draw, ((x0 + x1) / 2, y0 + max(3, h / 2 - 7)), money(item[key]), small_font, "#333333")
            current_bottom = y0
        draw_centered(draw, ((x0 + x1) / 2, plot_bottom + 14), item["label"], label_font)
    legend_y = 435
    for i, (label, color) in enumerate((("Series1", colors["dev"]), ("Series2", colors["qa"]), ("Series3", colors["prod"]))):
        x = 410 + i * 92
        draw.rectangle((x, legend_y, x + 10, legend_y + 10), fill=color)
        draw.text((x + 14, legend_y - 3), label, font=label_font, fill="#555555")
    img.save(ambiente_path)

    # Total monthly chart
    img = Image.new("RGB", (856, 640), "white")
    draw = ImageDraw.Draw(img)
    draw_centered(draw, (428, 24), f"Consumos totales por mes - IBM Cloud ({last_title})", title_font)
    plot_left, plot_top, plot_right, plot_bottom = 86, 80, 810, 520
    max_total = nice_axis_max(max(m["total"] for m in months))
    for tick in range(0, int(max_total) + 1, 500):
        if tick % 1000 == 0:
            y = plot_bottom - (tick / max_total) * (plot_bottom - plot_top)
            draw.line((plot_left, y, plot_right, y), fill="#D9D9D9", width=1)
            draw.text((6, y - 8), f"${tick:,.2f}", font=small_font, fill="#555555")
    gap = (plot_right - plot_left) / max(len(months), 1)
    bar_w = min(34, gap * 0.45)
    for idx, item in enumerate(months):
        x0 = plot_left + gap * idx + gap / 2 - bar_w / 2
        x1 = x0 + bar_w
        y0 = plot_bottom - (item["total"] / max_total) * (plot_bottom - plot_top)
        draw.rectangle((x0, y0, x1, plot_bottom), fill="#2F5597")
        draw_centered(draw, ((x0 + x1) / 2, plot_bottom + 14), item["label"], small_font)
    draw.rectangle((400, 585, 410, 595), fill="#2F5597")
    draw.text((416, 582), "TOTAL:", font=label_font, fill="#555555")
    img.save(total_path)

    return {"ambiente": ambiente_path, "total": total_path}


def collect_cuadro_data(comparativo_path: Path):
    with zipfile.ZipFile(comparativo_path, "r") as z:
        shared = workbook_shared_strings(z)
        sheet_path = workbook_sheet_path(z, "Cuadro")
        if sheet_path is None:
            return []
        rows = worksheet_row_maps(z, sheet_path, max_row=8)

    months = []
    for col, cell in sorted(rows.get(4, {}).items()):
        if col < 3:
            continue
        value = workbook_cell_value(cell, shared)
        month_date = excel_date_to_datetime(as_float(value)) if str(value or "").strip() else None
        if month_date:
            label = f"{MESES_ABREV[month_date.month]}-{str(month_date.year)[-2:]}"
        else:
            text = str(value or "").strip()
            if not re.fullmatch(r"[A-ZÁÉÍÓÚÑa-záéíóúñ]{3}-\d{2}", text):
                continue
            label = text.capitalize()
            year, month, _ = AUTOMATIZACION.parse_etiqueta_mes(label)
            month_date = datetime(year, month, 1)
        months.append(
            {
                "col": col,
                "date": month_date,
                "label": label,
                "dev": as_float(workbook_cell_value(rows.get(5, {}).get(col), shared)),
                "qa": as_float(workbook_cell_value(rows.get(6, {}).get(col), shared)),
                "prod": as_float(workbook_cell_value(rows.get(7, {}).get(col), shared)),
                "total": as_float(workbook_cell_value(rows.get(8, {}).get(col), shared)),
            }
        )
    deduped = {}
    for item in months:
        deduped[(item["date"].year, item["date"].month)] = item
    return sorted(deduped.values(), key=lambda item: (item["date"].year, item["date"].month))


def collect_resumen_month_data(comparativo_path: Path):
    with zipfile.ZipFile(comparativo_path, "r") as z:
        shared = workbook_shared_strings(z)
        sheet_path = workbook_sheet_path(z, "Resumen")
        if sheet_path is None:
            return []
        rows = worksheet_row_maps(z, sheet_path, max_row=120)

    total_general_row = None
    for row_idx, row in rows.items():
        if str(workbook_cell_value(row.get(1), shared) or "").strip().lower() == "total general":
            total_general_row = row_idx
            break
    if total_general_row is None:
        return []

    months = []
    row4 = rows.get(4, {})
    total_row = rows.get(total_general_row, {})
    for col, cell in sorted(row4.items()):
        label_text = str(workbook_cell_value(cell, shared) or "").strip()
        if not label_text.startswith("Total ") or label_text in {"Total general", "Total Var."}:
            continue
        label = label_text.replace("Total ", "", 1).strip()
        try:
            year, month, _ = AUTOMATIZACION.parse_etiqueta_mes(label)
        except ValueError:
            continue
        start_col = col - 6
        dev = as_float(workbook_cell_value(total_row.get(start_col), shared))
        devqa = as_float(workbook_cell_value(total_row.get(start_col + 1), shared))
        qa = as_float(workbook_cell_value(total_row.get(start_col + 2), shared))
        prod = as_float(workbook_cell_value(total_row.get(start_col + 4), shared))
        all_value = as_float(workbook_cell_value(total_row.get(start_col + 5), shared))
        months.append(
            {
                "col": col,
                "date": datetime(year, month, 1),
                "label": f"{MESES_ABREV[month]}-{str(year)[-2:]}",
                "dev": dev + (devqa / 2) + (all_value / 3),
                "qa": qa + (devqa / 2) + (all_value / 3),
                "prod": prod + (all_value / 3),
                "total": as_float(workbook_cell_value(total_row.get(col), shared)),
            }
        )

    deduped = {}
    for item in months:
        deduped[(item["date"].year, item["date"].month)] = item
    return sorted(deduped.values(), key=lambda item: (item["date"].year, item["date"].month))


def prepare_resource_images(image_paths: list[Path], output_dir: Path) -> list[Path]:
    if not image_paths:
        return []
    try:
        from PIL import Image, UnidentifiedImageError
    except ImportError as exc:
        raise ValueError("Falta instalar Pillow. Ejecuta: pip install -r requirements.txt") from exc

    prepared = []
    max_width = 1200
    max_height = 1200
    for index, path in enumerate(image_paths, 1):
        try:
            with Image.open(path) as opened:
                opened.load()
                image = opened.convert("RGB")
        except UnidentifiedImageError as exc:
            header = path.read_bytes()[:32].lower()
            if b"ftypavif" in header or b"ftypavis" in header:
                detail = "Parece ser AVIF aunque tenga extension JPG/JPEG."
            else:
                detail = "No se pudo reconocer como PNG/JPG/JPEG valido."
            raise ValueError(
                f"La imagen de consumo por recurso '{path.name}' no se puede procesar. "
                f"{detail} Abrela y exportala/guardala como PNG o JPG real, luego vuelve a cargarla."
            ) from exc
        except OSError as exc:
            raise ValueError(
                f"La imagen de consumo por recurso '{path.name}' no se puede leer. "
                "Guardala nuevamente como PNG o JPG e intenta otra vez."
            ) from exc
        ratio = min(max_width / image.width, max_height / image.height, 1)
        if ratio < 1:
            image = image.resize((int(image.width * ratio), int(image.height * ratio)))
        output = output_dir / f"consumo_por_recurso_{index:02d}.jpg"
        image.save(output, "JPEG", quality=82, optimize=True)
        prepared.append(output)
    return prepared


def collect_service_variations(comparativo_path: Path, threshold: float = 20.0, max_items: int = 6) -> list[str]:
    with zipfile.ZipFile(comparativo_path, "r") as z:
        shared = workbook_shared_strings(z)
        sheet_path = workbook_sheet_path(z, "Resumen")
        if sheet_path is None:
            return []
        rows = worksheet_row_maps(z, sheet_path, max_row=200)

    total_var_col = None
    total_general_row = None
    for col, cell in rows.get(4, {}).items():
        if str(workbook_cell_value(cell, shared) or "").strip() == "Total Var.":
            total_var_col = col
            break
    for row_idx, row in rows.items():
        if row_idx > 200:
            continue
        if str(workbook_cell_value(row.get(1), shared) or "").strip().lower() == "total general":
            total_general_row = row_idx
            break
    if not total_var_col or not total_general_row:
        return []

    env_cols = []
    for col in range(total_var_col - 6, total_var_col):
        label = str(workbook_cell_value(rows.get(5, {}).get(col), shared) or "").strip()
        if label:
            env_cols.append((label, col))

    variations = []
    for row in range(6, total_general_row):
        row_cells = rows.get(row, {})
        raw_service = str(workbook_cell_value(row_cells.get(1), shared) or "")
        service = raw_service.strip()
        if not service or service == "-":
            continue
        if raw_service != service:
            continue
        total_var = as_float(workbook_cell_value(row_cells.get(total_var_col), shared))
        if abs(total_var) < threshold:
            continue
        env_parts = []
        for env, col in env_cols:
            value = as_float(workbook_cell_value(row_cells.get(col), shared))
            if abs(value) >= 0.01:
                env_parts.append((env, value))
        env_parts.sort(key=lambda item: abs(item[1]), reverse=True)
        variations.append(
            {
                "service": service,
                "total": float(total_var),
                "envs": env_parts[:4],
            }
        )

    variations.sort(key=lambda item: abs(item["total"]), reverse=True)
    paragraphs = []
    for item in variations[:max_items]:
        service = item["service"]
        total = item["total"]
        amount = abs(total)
        if item["envs"]:
            env_text = ", ".join(f"{env}: {abs(value):,.2f} USD" for env, value in item["envs"])
            env_sentence = f" La variacion se concentra en {env_text}."
        else:
            env_sentence = ""
        if total > 0:
            paragraphs.append(
                f"En {service}, el costo aumento en {amount:,.2f} USD respecto al mes anterior.{env_sentence} "
                "Esta variacion representa un incremento de costos para el periodo y debe revisarse frente al comportamiento de uso, capacidad o configuracion del servicio."
            )
        else:
            paragraphs.append(
                f"En {service}, el costo disminuyo en {amount:,.2f} USD respecto al mes anterior.{env_sentence} "
                "Esta variacion es positiva para el periodo, al reflejar una reduccion del consumo facturado frente al mes previo."
            )

    if len(variations) > max_items:
        remaining = len(variations) - max_items
        paragraphs.append(
            f"Adicionalmente, se identifican {remaining} servicios con variaciones menores que no modifican de forma material la lectura general del consumo mensual."
        )
    return paragraphs


def money(value: float) -> str:
    return f"${value:,.2f}"


def replace_text_in_paragraph(paragraph, old: str, new: str):
    if old not in paragraph.text:
        return
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    if paragraph.runs:
        paragraph.runs[0].text = paragraph.text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""


def set_paragraph_text(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def update_word_table(document, months):
    if len(document.tables) < 2:
        return
    table = document.tables[1]
    needed_rows = len(months) + 2
    while len(table.rows) < needed_rows:
        table._tbl.append(copy.deepcopy(table.rows[-1]._tr))
    while len(table.rows) > needed_rows:
        row = table.rows[-1]._tr
        row.getparent().remove(row)

    first = months[0]["date"]
    last = months[-1]["date"]
    title = (
        "Consumos Totales Por Mes - IBM Cloud "
        f"({MESES_ABREV[first.month]} {first.year} - {MESES_ABREV[last.month]} {last.year})"
    )
    for cell in table.rows[0].cells:
        cell.text = title
    headers = ["Mes", "DEV", "QA", "PROD", "Consumo Total"]
    for idx, header in enumerate(headers):
        table.cell(1, idx).text = header
    for row_idx, month in enumerate(months, 2):
        table.cell(row_idx, 0).text = month["label"].lower()
        table.cell(row_idx, 1).text = money(month["dev"])
        table.cell(row_idx, 2).text = money(month["qa"])
        table.cell(row_idx, 3).text = money(month["prod"])
        table.cell(row_idx, 4).text = money(month["total"])


def update_word_text(document, months, etiqueta_mes: str):
    current = months[-1]
    previous = months[-2] if len(months) >= 2 else None
    parts = month_text_parts(etiqueta_mes)
    replacements = {
        "junio del 2026": f"{parts['name']} del {parts['year']}",
        "junio de 2026": f"{parts['name']} de {parts['year']}",
        "junio 2026": f"{parts['name']} {parts['year']}",
        "Junio 2026": f"{parts['name_cap']} {parts['year']}",
        "jun-26": str(parts["abbr_short"]).lower(),
        "Jun-26": str(parts["abbr_short"]),
        "Jun 2026": str(parts["title"]),
    }
    for paragraph in document.paragraphs:
        for old, new in replacements.items():
            replace_text_in_paragraph(paragraph, old, new)
        text = paragraph.text.strip()
        if text.startswith("En esta secciÃ³n se describe cada uno de los meses reportados"):
            set_paragraph_text(
                paragraph,
                "En esta secciÃ³n se describe cada uno de los meses reportados "
                f"hasta el mes de {parts['name']} de {parts['year']} en el orden de su ejecuciÃ³n. "
                "A continuaciÃ³n, se muestra una tabla resumen:",
            )
        elif text.startswith("En el mes de "):
            set_paragraph_text(
                paragraph,
                f"En el mes de {parts['name']} de {parts['year']}, el gasto total registrado en IBM Cloud "
                f"fue de USD. {current['total']:,.2f}. Durante este mes, el ambiente de Calidad, "
                "Desarrollo y ProducciÃ³n mantuvo la continuidad de sus actividades ya iniciadas previamente.",
            )
        elif text.startswith("En el presente mes se registra"):
            if previous:
                diff = current["total"] - previous["total"]
                pct = abs(diff) / previous["total"] * 100 if previous["total"] else 0
                prev_parts = month_text_parts(previous["label"])
                if abs(diff) < 0.005:
                    set_paragraph_text(
                        paragraph,
                        f"En el presente mes se mantiene el gasto total respecto al mes anterior, pasando de "
                        f"{previous['total']:,.2f} USD en {prev_parts['name']} a {current['total']:,.2f} USD "
                        f"en {parts['name']}, lo que representa una variaciÃ³n de cerca del {pct:.2f}%.",
                    )
                else:
                    direction = "un incremento" if diff > 0 else "una reducciÃ³n"
                    qualification = "negativa" if diff > 0 else "positiva"
                    set_paragraph_text(
                        paragraph,
                        f"En el presente mes se registra {direction} total de aproximadamente {abs(diff):,.2f} USD "
                        f"respecto al mes anterior, pasando de {previous['total']:,.2f} USD en {prev_parts['name']} "
                        f"a {current['total']:,.2f} USD en {parts['name']}, lo que representa una variaciÃ³n "
                        f"{qualification} de cerca del {pct:.2f}%.",
                    )


def format_report_file_references(document):
    targets = ("Instances (", "Summary (")
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text.startswith(targets):
            continue
        for run in paragraph.runs:
            run.text = ""
            run.bold = False
        if "):" in text and not text.endswith(".xlsx"):
            marker = text.index("):") + 2
            label = text[:marker]
            description = text[marker:]
            bold_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            bold_run.text = label
            bold_run.bold = True
            normal_run = paragraph.add_run(description)
            normal_run.bold = False
        else:
            normal_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            normal_run.text = text
            normal_run.bold = False


def patch_docx_media(docx_path: Path, replacements: dict[str, Path]):
    temp_docx = docx_path.with_name(f"{docx_path.stem}_media{docx_path.suffix}")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(temp_docx, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename in replacements:
                dst.writestr(item, replacements[item.filename].read_bytes())
            else:
                dst.writestr(item, src.read(item.filename))
    temp_docx.replace(docx_path)


def generar_informe_word(
    template_path: Path,
    salida_word: Path,
    comparativo_path: Path,
    etiqueta_mes: str,
    run_output_dir: Path,
    resource_images: list[Path],
):
    started = time.perf_counter()
    print("[Word] Iniciando informe...", flush=True)
    step_started = time.perf_counter()
    months = collect_resumen_month_data(comparativo_path) or collect_cuadro_data(comparativo_path)
    if not months:
        raise ValueError("No se encontraron meses en el comparativo.")
    print(f"[Word] Meses leidos para informe: {len(months)} ({time.perf_counter() - step_started:.1f}s)", flush=True)
    step_started = time.perf_counter()
    exported = chart_export_paths(comparativo_path, run_output_dir, months)
    print(f"[Word] Graficos listos. ({time.perf_counter() - step_started:.1f}s)", flush=True)
    step_started = time.perf_counter()
    prepared_resource_images = prepare_resource_images(resource_images, run_output_dir)
    print(f"[Word] Imagenes de consumo por recurso listas: {len(prepared_resource_images)} ({time.perf_counter() - step_started:.1f}s)", flush=True)
    step_started = time.perf_counter()
    service_variations = collect_service_variations(comparativo_path)
    print(f"[Word] Variaciones detectadas: {len(service_variations)} ({time.perf_counter() - step_started:.1f}s)", flush=True)

    try:
        from docx import Document
        from docx.shared import Cm
        from docxtpl import DocxTemplate, InlineImage
    except ImportError as exc:
        raise ValueError("Falta instalar docxtpl/python-docx. Ejecuta: pip install -r requirements.txt") from exc

    current = months[-1]
    previous = months[-2] if len(months) >= 2 else None
    parts = month_text_parts(etiqueta_mes)
    first = months[0]["date"]
    last = months[-1]["date"]
    context = {
        "mes_nombre": parts["name"],
        "mes_nombre_cap": parts["name_cap"],
        "mes_nombre_mayus": str(parts["name"]).upper(),
        "anio": parts["year"],
        "mes_abrev": parts["abbr_short"],
        "periodo_inicio": f"{MESES_ABREV[first.month]} {first.year}",
        "periodo_fin": f"{MESES_ABREV[last.month]} {last.year}",
        "periodo_inicio_largo": f"{MESES_LARGO[first.month]} de {first.year}",
        "total_mes": f"{current['total']:,.2f}",
        "fecha_cambio": datetime.now().strftime("%d/%m/%Y"),
        "detalle_variaciones": "",
    }
    if previous:
        diff = current["total"] - previous["total"]
        pct = abs(diff) / previous["total"] * 100 if previous["total"] else 0
        prev_parts = month_text_parts(previous["label"])
        context.update(
            {
                "mes_anterior_nombre": prev_parts["name"],
                "total_mes_anterior": f"{previous['total']:,.2f}",
                "variacion_monto": f"{abs(diff):,.2f}",
                "variacion_porcentaje": f"{pct:.2f}",
            }
        )
        if abs(diff) < 0.005:
            context["parrafo_variacion_general"] = (
                f"En el presente mes se mantiene el gasto total respecto al mes anterior, pasando de "
                f"{previous['total']:,.2f} USD en {prev_parts['name']} a {current['total']:,.2f} USD "
                f"en {parts['name']}, lo que representa una variacion de cerca del {pct:.2f}%."
            )
            context["detalle_variaciones"] = (
                "No se identifican incrementos o decrementos relevantes por servicio respecto al mes anterior. "
                "Los consumos se mantienen estables, por lo que no se evidencian cambios de configuracion, "
                "capacidad o uso que requieran una explicacion adicional en este periodo."
            )
        else:
            direction = "un incremento" if diff > 0 else "una reduccion"
            qualification = "negativa" if diff > 0 else "positiva"
            context["parrafo_variacion_general"] = (
                f"En el presente mes se registra {direction} total de aproximadamente {abs(diff):,.2f} USD "
                f"respecto al mes anterior, pasando de {previous['total']:,.2f} USD en {prev_parts['name']} "
                f"a {current['total']:,.2f} USD en {parts['name']}, lo que representa una variacion "
                f"{qualification} de cerca del {pct:.2f}%."
            )
            if service_variations:
                context["detalle_variaciones"] = "Durante este mes se registran las siguientes variaciones a nivel del consumo de creditos:\n\n" + "\n".join(
                    f"- {paragraph}" for paragraph in service_variations
                )
            else:
                context["detalle_variaciones"] = (
                    "Se identifican variaciones frente al mes anterior; sin embargo, no se detectan servicios individuales "
                    "que superen el umbral definido para una explicacion detallada."
                )
    else:
        context["parrafo_variacion_general"] = "No se cuenta con un mes anterior para calcular la variacion mensual."

    print("[Word] Renderizando plantilla...", flush=True)
    tpl = DocxTemplate(template_path)
    context["grafico_consumos_ambiente"] = InlineImage(tpl, str(exported["ambiente"]), width=Cm(15.37), height=Cm(9.59))
    context["grafico_consumos_totales"] = InlineImage(tpl, str(exported["total"]), width=Cm(15.37), height=Cm(9.59))
    context["imagen_consumo_recurso_1"] = (
        InlineImage(tpl, str(prepared_resource_images[0]), width=Cm(15.37), height=Cm(4.0))
        if len(prepared_resource_images) >= 1
        else ""
    )
    context["imagen_consumo_recurso_2"] = (
        InlineImage(tpl, str(prepared_resource_images[1]), width=Cm(15.32), height=Cm(10.95))
        if len(prepared_resource_images) >= 2
        else ""
    )
    context["imagen_consumo_recurso"] = context["imagen_consumo_recurso_1"]
    tpl.render(context)
    tpl.save(salida_word)
    print("[Word] Plantilla guardada, aplicando ajustes finales...", flush=True)

    document = Document(salida_word)
    update_word_table(document, months)
    format_report_file_references(document)
    document.save(salida_word)
    elapsed = time.perf_counter() - started
    print(f"[Word] Informe generado en {elapsed:.1f}s: {salida_word.name}", flush=True)


FRESHDESK_STATUS = {
    2: "Abierto",
    3: "Pendiente",
    4: "Resuelto",
    5: "Cerrado",
    6: "Esperando cliente",
    7: "Esperando tercero",
}


def parse_iso_datetime(value: str | None):
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None


def format_ticket_date(value: str | None):
    dt = parse_iso_datetime(value)
    if not dt:
        return "-"
    return f"{dt.day} {MESES_LARGO[dt.month]}"


def month_title_from_date(value: datetime):
    return f"{MESES_LARGO[value.month].capitalize()} {value.year}"


def freshdesk_credentials():
    api_key = os.getenv("FRESHDESK_API_KEY", "").strip()
    domain = os.getenv("FRESHDESK_DOMAIN", "").strip()
    domain = domain.replace("https://", "").replace("http://", "").strip("/")
    if not api_key or not domain:
        raise ValueError("Configura FRESHDESK_API_KEY y FRESHDESK_DOMAIN en el archivo .env.")
    return api_key, domain


def fetch_freshdesk_tickets(fecha_inicio: str, fecha_fin: str, responder_id: str = ""):
    try:
        import requests
    except ImportError as exc:
        raise ValueError("Falta instalar requests/python-dotenv. Ejecuta: pip install -r requirements.txt") from exc

    api_key, domain = freshdesk_credentials()
    start = datetime.strptime(fecha_inicio, "%Y-%m-%d")
    end = datetime.strptime(fecha_fin, "%Y-%m-%d")
    if end < start:
        raise ValueError("La fecha fin no puede ser menor que la fecha inicio.")

    # Freshdesk usa comparadores estrictos; se consulta hasta el dia siguiente para cubrir el cierre del rango.
    end_exclusive = (end + timedelta(days=1)).strftime("%Y-%m-%d")
    base_query = f"created_at:>'{fecha_inicio}' AND created_at:<'{end_exclusive}'"
    responder_id = responder_id.strip()
    query = f"{base_query} AND responder_id:{responder_id}" if responder_id else base_query
    url = f"https://{domain}/api/v2/search/tickets"

    def request_search(search_query: str):
        found: list[dict] = []
        total = None
        page = 1
        while True:
            response = requests.get(
                url,
                auth=(api_key, "X"),
                params={"query": f'"{search_query}"', "page": page},
                timeout=30,
            )
            if response.status_code >= 400:
                raise ValueError(f"Freshdesk respondio {response.status_code}: {response.text[:400]}")
            payload = response.json()
            if total is None:
                total = int(payload.get("total", 0))
            results = payload.get("results", [])
            found.extend(results)
            if not results or len(found) >= total:
                break
            page += 1
        return found

    try:
        tickets = request_search(query)
    except ValueError:
        if not responder_id:
            raise
        tickets = request_search(base_query)

    if responder_id:
        tickets = [ticket for ticket in tickets if str(ticket.get("responder_id") or "").strip() == responder_id]
    return tickets, query


def summarize_tickets(tickets: list[dict]):
    closed_statuses = {4, 5}
    incidents = [t for t in tickets if str(t.get("type") or "").strip().lower() == "incidente"]
    closed = [t for t in tickets if int(t.get("status") or 0) in closed_statuses]
    pending = [t for t in tickets if int(t.get("status") or 0) not in closed_statuses]
    response_values = []
    incident_resolution_values = []

    for ticket in tickets:
        created = parse_iso_datetime(ticket.get("created_at"))
        first_response = parse_iso_datetime(ticket.get("first_responded_at"))
        if created and first_response:
            response_values.append((first_response - created).total_seconds() / 60)

    for ticket in incidents:
        created = parse_iso_datetime(ticket.get("created_at"))
        closed_at = parse_iso_datetime(ticket.get("closed_at") or ticket.get("resolved_at"))
        if created and closed_at:
            incident_resolution_values.append((closed_at - created).total_seconds() / 3600)

    response_ok = sum(1 for minutes in response_values if minutes <= 30)
    incident_ok = sum(1 for hours in incident_resolution_values if hours <= 4)
    return {
        "total": len(tickets),
        "closed": len(closed),
        "pending": len(pending),
        "incidents": len(incidents),
        "response_measured": len(response_values),
        "response_ok": response_ok,
        "response_pct": (response_ok / len(response_values) * 100) if response_values else None,
        "incident_measured": len(incident_resolution_values),
        "incident_ok": incident_ok,
        "incident_pct": (incident_ok / len(incident_resolution_values) * 100) if incident_resolution_values else None,
    }


def is_helpdesk_ticket_table(table):
    return bool(table.rows and table.rows[0].cells and "ID Ticket" in table.rows[0].cells[0].text)


def iter_document_paragraphs(document, skip_ticket_table: bool = False):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        if skip_ticket_table and is_helpdesk_ticket_table(table):
            continue
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_helpdesk_text(document, replacements: dict[str, str], skip_ticket_table: bool = False):
    for paragraph in iter_document_paragraphs(document, skip_ticket_table=skip_ticket_table):
        for old, new in replacements.items():
            replace_text_in_paragraph(paragraph, old, new)


def ticket_id_from_text(value: str):
    match = re.search(r"\d+", value or "")
    return match.group(0) if match else ""


def ticket_numeric_id(ticket: dict):
    try:
        return int(ticket.get("id") or 0)
    except (TypeError, ValueError):
        return 0


def ticket_row_values(ticket: dict):
    status = FRESHDESK_STATUS.get(int(ticket.get("status") or 0), str(ticket.get("status") or "-"))
    return [
        f"N° {ticket.get('id', '-')}",
        str(ticket.get("type") or "Solicitud"),
        str(ticket.get("subject") or "-"),
        status,
        format_ticket_date(ticket.get("created_at")),
        format_ticket_date(ticket.get("closed_at") or ticket.get("resolved_at")),
        "Manuel Alcalá",
    ]


def set_cell_text_preserve_style(cell, value: str):
    if not cell.paragraphs:
        cell.text = value
        return cell.paragraphs[0].runs[0] if cell.paragraphs and cell.paragraphs[0].runs else None
    first_paragraph = cell.paragraphs[0]
    if first_paragraph.runs:
        first_paragraph.runs[0].text = value
        primary_run = first_paragraph.runs[0]
        for run in first_paragraph.runs[1:]:
            run.text = ""
    else:
        primary_run = first_paragraph.add_run(value)
    for paragraph in cell.paragraphs[1:]:
        paragraph.clear()
    return primary_run


def set_row_values_preserve_style(row, values: list[str]):
    from docx.enum.text import WD_COLOR_INDEX

    for idx, value in enumerate(values):
        if idx < len(row.cells):
            run = set_cell_text_preserve_style(row.cells[idx], value)
            if run is not None and idx == 3:
                status = value.strip().lower()
                if status == "pendiente":
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif status == "cerrado":
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                else:
                    run.font.highlight_color = None


def sort_ticket_table_rows_desc(ticket_table):
    data_rows = list(ticket_table.rows[1:])
    data_rows.sort(key=lambda row: int(ticket_id_from_text(row.cells[0].text) or 0), reverse=True)
    for row in data_rows:
        ticket_table._tbl.append(row._tr)


def apply_ticket_status_highlights(ticket_table):
    from docx.enum.text import WD_COLOR_INDEX

    for row in ticket_table.rows[1:]:
        if len(row.cells) < 4:
            continue
        status_cell = row.cells[3]
        status_text = status_cell.text.strip().lower()
        for paragraph in status_cell.paragraphs:
            for run in paragraph.runs:
                if status_text == "pendiente":
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif status_text == "cerrado":
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                else:
                    run.font.highlight_color = None


def update_helpdesk_ticket_table(document, tickets: list[dict]):
    ticket_table = None
    for table in document.tables:
        if is_helpdesk_ticket_table(table):
            ticket_table = table
            break
    if ticket_table is None:
        return []

    existing_ids = set()
    placeholder_rows = []
    for row in ticket_table.rows[1:]:
        row_id = ticket_id_from_text(row.cells[0].text if row.cells else "")
        if row_id:
            existing_ids.add(row_id)
        elif any("Sin tickets" in cell.text for cell in row.cells):
            placeholder_rows.append(row)

    missing_tickets = [ticket for ticket in tickets if str(ticket.get("id") or "") not in existing_ids]
    missing_tickets.sort(key=ticket_numeric_id, reverse=True)

    if not missing_tickets and len(ticket_table.rows) == 1:
        row = ticket_table.add_row()
        values = ["-", "-", "Sin tickets en el rango seleccionado", "-", "-", "-", "-"]
        for idx, value in enumerate(values):
            row.cells[idx].text = value
        return []

    if missing_tickets and placeholder_rows:
        for row in placeholder_rows:
            row._tr.getparent().remove(row._tr)

    style_source_row = ticket_table.rows[1] if len(ticket_table.rows) > 1 else None
    for ticket in missing_tickets:
        if style_source_row is not None:
            new_tr = copy.deepcopy(style_source_row._tr)
            ticket_table.rows[0]._tr.addnext(new_tr)
            row = ticket_table.rows[1]
        else:
            row = ticket_table.add_row()
        values = ticket_row_values(ticket)
        set_row_values_preserve_style(row, values)
    sort_ticket_table_rows_desc(ticket_table)
    apply_ticket_status_highlights(ticket_table)
    return missing_tickets


def find_section_paragraph(document, text: str):
    needle = text.lower()
    for paragraph in document.paragraphs:
        if needle in paragraph.text.lower():
            return paragraph
    return None


def paragraph_text_from_element(element):
    return "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t"))


def paragraph_from_element(document, element):
    from docx.text.paragraph import Paragraph

    return Paragraph(element, document)


def find_ticket_table(document):
    for table in document.tables:
        if is_helpdesk_ticket_table(table):
            return table
    return None


def find_first_case_after_ticket_table(document):
    ticket_table = find_ticket_table(document)
    if ticket_table is None:
        return None

    found_ticket_table = False
    for element in document.element.body:
        if element is ticket_table._tbl:
            found_ticket_table = True
            continue
        if not found_ticket_table or not element.tag.endswith("}p"):
            continue
        text = paragraph_text_from_element(element).strip()
        if re.match(r"^[a-z]\.\s*Caso\s+\d+", text, flags=re.IGNORECASE) or text.startswith("Caso "):
            return paragraph_from_element(document, element)
        if "Etiquetado de componentes" in text:
            return paragraph_from_element(document, element)
    return None


def find_case_insert_anchor(document):
    table_index = ticket_table_body_index(document)
    if table_index < 0:
        return find_section_paragraph(document, "Etiquetado de componentes")

    elements = list(document.element.body)
    section_idx = None
    last_case_element = None
    for idx in range(table_index + 1, len(elements)):
        element = elements[idx]
        if not element.tag.endswith("}p"):
            continue
        text = paragraph_text_from_element(element).strip()
        if "Etiquetado de componentes" in text:
            section_idx = idx
            break
        if re.match(r"^[a-z]\.\s*Caso\s+\d+", text, flags=re.IGNORECASE) or text.startswith("Caso "):
            last_case_element = element

    if section_idx is not None:
        return paragraph_from_element(document, elements[section_idx])
    if last_case_element is not None:
        return paragraph_from_element(document, last_case_element)
    return find_first_case_after_ticket_table(document)


def find_case_insert_context(document):
    table_index = ticket_table_body_index(document)
    if table_index < 0:
        anchor = find_section_paragraph(document, "Etiquetado de componentes")
        return anchor, anchor

    elements = list(document.element.body)
    section_element = None
    last_case_element = None
    for idx in range(table_index + 1, len(elements)):
        element = elements[idx]
        if not element.tag.endswith("}p"):
            continue
        text = paragraph_text_from_element(element).strip()
        if "Etiquetado de componentes" in text:
            section_element = element
            break
        if re.match(r"^[a-z]\.\s*Caso\s+\d+", text, flags=re.IGNORECASE) or text.startswith("Caso "):
            last_case_element = element

    anchor = paragraph_from_element(document, section_element) if section_element is not None else find_case_insert_anchor(document)
    template = paragraph_from_element(document, last_case_element) if last_case_element is not None else anchor
    return anchor, template


def ticket_table_body_index(document):
    ticket_table = find_ticket_table(document)
    if ticket_table is None:
        return -1
    for idx, element in enumerate(document.element.body):
        if element is ticket_table._tbl:
            return idx
    return -1


def paragraph_index(document, target):
    for idx, paragraph in enumerate(document.paragraphs):
        if paragraph._p is target._p:
            return idx
    return -1


def paragraph_after(document, paragraph):
    paragraphs = document.paragraphs
    index = paragraph_index(document, paragraph)
    if 0 <= index + 1 < len(paragraphs):
        return paragraphs[index + 1]
    return None


def copy_paragraph_properties(source, target):
    if source is None or source._p.pPr is None:
        return
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    target._p.insert(0, copy.deepcopy(source._p.pPr))


def helpdesk_case_templates(document, anchor):
    title_template = anchor
    detail_template = paragraph_after(document, anchor)
    return title_template, detail_template


def add_ticket_detail_before_section(document, tickets: list[dict]):
    if not tickets:
        return
    anchor, title_template_anchor = find_case_insert_context(document)
    if anchor is None:
        return
    title_template, detail_template = helpdesk_case_templates(document, title_template_anchor)

    for ticket in sorted(tickets, key=ticket_numeric_id):
        title = anchor.insert_paragraph_before()
        copy_paragraph_properties(title_template, title)
        title_run = title.add_run(f"Caso {ticket.get('id', '-')}: {ticket.get('subject') or 'Ticket Freshdesk'}")
        title_run.bold = True

        request = anchor.insert_paragraph_before()
        copy_paragraph_properties(detail_template, request)
        request.add_run("Solicitud:")

        analysis = anchor.insert_paragraph_before()
        copy_paragraph_properties(detail_template, analysis)
        analysis.add_run("Análisis:")

        result = anchor.insert_paragraph_before()
        copy_paragraph_properties(detail_template, result)
        result.add_run("Resultado:")

        blank = anchor.insert_paragraph_before()
        copy_paragraph_properties(detail_template, blank)


def existing_case_ids(document):
    found = set()
    table_index = ticket_table_body_index(document)
    for idx, element in enumerate(document.element.body):
        if table_index >= 0 and idx <= table_index:
            continue
        if not element.tag.endswith("}p"):
            continue
        match = re.search(r"\bCaso\s+(\d+)\b", paragraph_text_from_element(element), flags=re.IGNORECASE)
        if match:
            found.add(match.group(1))
    return found


def remove_misplaced_ticket_details(document, tickets: list[dict]):
    target_ids = {str(ticket.get("id") or "") for ticket in tickets}
    target_ids.discard("")
    if not target_ids:
        return
    table_index = ticket_table_body_index(document)
    if table_index < 0:
        return

    body = document.element.body
    elements = list(body)
    remove_elements = []
    idx = 0
    while idx < table_index:
        element = elements[idx]
        text = paragraph_text_from_element(element).strip() if element.tag.endswith("}p") else ""
        match = re.search(r"\bCaso\s+(\d+)\b", text, flags=re.IGNORECASE)
        if not match or match.group(1) not in target_ids:
            idx += 1
            continue

        remove_elements.append(element)
        idx += 1
        while idx < table_index:
            next_element = elements[idx]
            if not next_element.tag.endswith("}p"):
                break
            next_text = paragraph_text_from_element(next_element).strip()
            remove_elements.append(next_element)
            idx += 1
            if not next_text or next_text.lower().startswith("resultado:"):
                if idx < table_index:
                    maybe_blank = elements[idx]
                    if maybe_blank.tag.endswith("}p") and not paragraph_text_from_element(maybe_blank).strip():
                        remove_elements.append(maybe_blank)
                        idx += 1
                break

    for element in remove_elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def remove_empty_ticket_details_after_table(document, tickets: list[dict]):
    target_ids = {str(ticket.get("id") or "") for ticket in tickets}
    target_ids.discard("")
    table_index = ticket_table_body_index(document)
    if table_index < 0 or not target_ids:
        return

    body = document.element.body
    elements = list(body)
    remove_elements = []
    idx = table_index + 1
    while idx < len(elements):
        element = elements[idx]
        if element.tag.endswith("}p") and "Etiquetado de componentes" in paragraph_text_from_element(element):
            break
        text = paragraph_text_from_element(element).strip() if element.tag.endswith("}p") else ""
        match = re.search(r"\bCaso\s+(\d+)\b", text, flags=re.IGNORECASE)
        if not match or match.group(1) not in target_ids:
            idx += 1
            continue

        block = [element]
        probe = idx + 1
        labels = []
        while probe < len(elements):
            next_element = elements[probe]
            if not next_element.tag.endswith("}p"):
                break
            next_text = paragraph_text_from_element(next_element).strip()
            if re.search(r"\bCaso\s+\d+\b", next_text, flags=re.IGNORECASE) or "Etiquetado de componentes" in next_text:
                break
            block.append(next_element)
            if next_text:
                labels.append(next_text.lower())
            probe += 1
            if next_text.lower().startswith("resultado:"):
                if probe < len(elements):
                    maybe_blank = elements[probe]
                    if maybe_blank.tag.endswith("}p") and not paragraph_text_from_element(maybe_blank).strip():
                        block.append(maybe_blank)
                        probe += 1
                break

        label_set = {re.sub(r"\s+", " ", label) for label in labels}
        if {"solicitud:", "análisis:", "resultado:"}.issubset(label_set) or {"solicitud:", "analisis:", "resultado:"}.issubset(label_set):
            remove_elements.extend(block)
        idx = max(probe, idx + 1)

    for element in remove_elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def tickets_missing_case_detail(document, tickets: list[dict]):
    existing = existing_case_ids(document)
    return [ticket for ticket in tickets if str(ticket.get("id") or "") not in existing]


def update_helpdesk_version_date(document, generated_at: datetime):
    date_value = generated_at.strftime("%d/%m/%Y")
    for table in document.tables:
        if not table.rows:
            continue
        header_cells = table.rows[0].cells
        date_col = None
        for idx, cell in enumerate(header_cells):
            normalized = re.sub(r"\s+", " ", cell.text.upper())
            if "FECHA" in normalized and "CAMBIO" in normalized:
                date_col = idx
                break
        if date_col is None or len(table.rows) < 2:
            continue
        set_cell_text_preserve_style(table.rows[1].cells[date_col], date_value)
        return


def generar_informe_helpdesk(
    template_path: Path,
    salida_word: Path,
    tickets: list[dict],
    fecha_inicio: str,
    fecha_fin: str,
):
    if not template_path.exists():
        raise ValueError(f"No se encontro la plantilla Helpdesk: {template_path}")
    from docx import Document

    start = datetime.strptime(fecha_inicio, "%Y-%m-%d")
    end = datetime.strptime(fecha_fin, "%Y-%m-%d")
    summary = summarize_tickets(tickets)
    month_title = month_title_from_date(end)
    month_name = MESES_LARGO[end.month]
    cierre = f"{end.day} {month_name}"
    response_text = (
        f"{summary['response_ok']} de {summary['response_measured']} = {summary['response_pct']:.0f}%"
        if summary["response_pct"] is not None
        else "No disponible con los campos retornados por Freshdesk"
    )
    incident_text = (
        f"{summary['incident_ok']} de {summary['incident_measured']} = {summary['incident_pct']:.0f}%"
        if summary["incident_pct"] is not None
        else "No disponible con los campos retornados por Freshdesk"
    )

    document = Document(template_path)
    replace_helpdesk_text(
        document,
        {
            "Mayo 2026": month_title,
            "Junio:": f"{month_name.capitalize()}:",
            "30 junio": cierre,
            "mes de junio": f"mes de {month_name}",
            "junio": month_name,
            "5 tickets": f"{summary['total']} tickets",
            "9 ticket": f"{summary['total']} ticket",
            "2 cerrados en el presente mes y 2 por atender": f"{summary['closed']} cerrados y {summary['pending']} pendientes",
            "5 5 =100%": response_text,
            "1 1 =100%": incident_text,
            "Durante el mes solo se presento 1 caso de incidencia": (
                f"Durante el periodo se registraron {summary['incidents']} caso(s) de incidencia."
            ),
        },
        skip_ticket_table=True,
    )
    added_tickets = update_helpdesk_ticket_table(document, tickets)
    remove_misplaced_ticket_details(document, tickets)
    remove_empty_ticket_details_after_table(document, tickets)
    detail_tickets = tickets_missing_case_detail(document, tickets)
    add_ticket_detail_before_section(document, detail_tickets)
    update_helpdesk_version_date(document, datetime.now())
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("A la fecha de cierre"):
            set_paragraph_text(
                paragraph,
                f"A la fecha de cierre ({cierre}) se han generado {summary['total']} tickets en el periodo "
                f"seleccionado. Del total indicado, {summary['closed']} se encuentran cerrados o resueltos y "
                f"{summary['pending']} permanecen pendientes o en atencion.",
            )
        elif text.startswith("De los ") and "tickets" in text and ("presente mes" in text or "periodo" in text):
            set_paragraph_text(
                paragraph,
                f"De los {summary['total']} tickets del periodo se han cerrado {summary['closed']}, "
                f"quedando {summary['pending']} pendientes de atencion. El detalle se muestra en el siguiente cuadro.",
            )

    document.save(salida_word)
    return len(added_tickets)


def run_helpdesk_automation(fields: dict[str, str], files: dict[str, tuple[str, bytes]] | None = None):
    started = time.perf_counter()
    run_id = uuid.uuid4().hex[:12]
    run_output_dir = OUTPUT_DIR / run_id
    run_output_dir.mkdir(parents=True, exist_ok=True)
    fecha_inicio = fields.get("fecha_inicio", "").strip()
    fecha_fin = fields.get("fecha_fin", "").strip()
    responder_id = os.getenv("FRESHDESK_RESPONDER_ID", "").strip()
    if not fecha_inicio or not fecha_fin:
        raise ValueError("Ingresa fecha inicio y fecha fin.")
    if not responder_id:
        raise ValueError("Configura FRESHDESK_RESPONDER_ID en el archivo .env.")
    print("[Helpdesk] Consultando Freshdesk...", flush=True)
    tickets, query = fetch_freshdesk_tickets(fecha_inicio, fecha_fin, responder_id)
    end = datetime.strptime(fecha_fin, "%Y-%m-%d")
    salida_word = run_output_dir / f"XM-Soporte-Informe-Helpdesk-{MESES_ABREV[end.month]}-{end.year}.docx"
    template_path = HELPDESK_TEMPLATE_PATH
    if files and "helpdesk_base" in files:
        filename, content = files["helpdesk_base"]
        template_path = run_output_dir / safe_filename(filename, ".docx")
        template_path.write_bytes(content)
    print(f"[Helpdesk] Tickets recibidos: {len(tickets)}. Generando Word...", flush=True)
    added_count = generar_informe_helpdesk(
        template_path,
        salida_word,
        tickets,
        fecha_inicio,
        fecha_fin,
    )
    print(f"[Helpdesk] Informe generado en {time.perf_counter() - started:.1f}s", flush=True)
    return {
        "files": [{"label": salida_word.name, "url": f"/download?run={run_id}&file={salida_word.name}"}],
        "count": len(tickets),
        "added_count": added_count,
        "responder_id": responder_id,
        "query": query,
    }


def run_automation_v2(fields: dict[str, str], files: dict[str, tuple[str, bytes]]):
    started = time.perf_counter()
    run_id = uuid.uuid4().hex[:12]
    run_upload_dir = UPLOAD_DIR / run_id
    run_output_dir = OUTPUT_DIR / run_id
    run_upload_dir.mkdir(parents=True, exist_ok=True)
    run_output_dir.mkdir(parents=True, exist_ok=True)

    try:
        print("[Run] Recibiendo archivos...", flush=True)
        step_started = time.perf_counter()
        summary_path = save_upload(files, "summary", run_upload_dir, required=True)
        instances_path = save_upload(files, "instances", run_upload_dir, required=True)
        comparativo_path = save_upload(files, "comparativo", run_upload_dir, required=False)
        resource_images = save_uploads(files, "resource_images", run_upload_dir, extensions=(".png", ".jpg", ".jpeg"))
        mes = resolve_month(fields, files)
        print(f"[Run] Archivos listos para {mes}. ({time.perf_counter() - step_started:.1f}s)", flush=True)
        duplicate_action = (fields.get("duplicate_action") or "").strip()

        step_started = time.perf_counter()
        print("[Run] Validando mes duplicado...", flush=True)
        if comparativo_path is not None and comparativo_has_month(comparativo_path, mes) and duplicate_action not in {"agregar", "actualizar"}:
            print(f"[Run] Mes duplicado detectado. ({time.perf_counter() - step_started:.1f}s)", flush=True)
            return {
                "needs_confirmation": True,
                "month": mes,
                "message": f"El comparativo ya contiene {mes}.",
            }
        print(f"[Run] Validacion de duplicado lista. ({time.perf_counter() - step_started:.1f}s)", flush=True)

        step_started = time.perf_counter()
        print("[Run] Consolidando XM Summary...", flush=True)
        summary = AUTOMATIZACION.consolidar_summary(summary_path)
        print(f"[Run] XM Summary consolidado. ({time.perf_counter() - step_started:.1f}s)", flush=True)
        step_started = time.perf_counter()
        print("[Run] Consolidando XM Instances...", flush=True)
        por_plan, por_servicio, servicios_planes = AUTOMATIZACION.consolidar_instances(instances_path)
        print(f"[Run] XM Instances consolidado. ({time.perf_counter() - step_started:.1f}s)", flush=True)

        year_month, month_name = month_filename_parts(mes)
        salida_summary = run_output_dir / f"XM-Summary-{year_month}.xlsx"
        salida_instances = run_output_dir / f"XM-Instances-{year_month}.xlsx"
        step_started = time.perf_counter()
        print("[Run] Generando archivos XM separados...", flush=True)
        AUTOMATIZACION.generar_excels_tablas(
            salida_summary,
            salida_instances,
            summary,
            por_plan,
            por_servicio,
            servicios_planes,
        )
        print(f"[Run] Archivos XM generados. ({time.perf_counter() - step_started:.1f}s)", flush=True)

        outputs = [
            {"label": salida_summary.name, "url": f"/download?run={run_id}&file={salida_summary.name}"},
            {"label": salida_instances.name, "url": f"/download?run={run_id}&file={salida_instances.name}"},
        ]

        if comparativo_path is not None:
            first_month = comparativo_first_month(comparativo_path) or month_name
            salida_comparativo = run_output_dir / f"Comparativo-Instances-{first_month}-{month_name}.xlsx"
            step_started = time.perf_counter()
            print("[Run] Actualizando comparativo...", flush=True)
            AUTOMATIZACION.actualizar_comparativo_xml(
                comparativo_path,
                salida_comparativo,
                mes,
                por_plan,
                por_servicio,
                duplicate_action or "auto",
                instances_path,
            )
            print(f"[Run] Comparativo actualizado. ({time.perf_counter() - step_started:.1f}s)", flush=True)
            outputs.append(
                {
                    "label": salida_comparativo.name,
                    "url": f"/download?run={run_id}&file={salida_comparativo.name}",
                }
            )

            if not REPORT_TEMPLATE_PATH.exists():
                raise ValueError(f"No se encontro la plantilla Word: {REPORT_TEMPLATE_PATH}")

            _, report_month = month_filename_parts(mes)
            salida_word = run_output_dir / f"Informe-de-Consumos-XM-{report_month}.docx"
            step_started = time.perf_counter()
            print("[Run] Generando informe Word...", flush=True)
            generar_informe_word(
                REPORT_TEMPLATE_PATH,
                salida_word,
                salida_comparativo,
                mes,
                run_output_dir,
                resource_images,
            )
            print(f"[Run] Informe Word generado. ({time.perf_counter() - step_started:.1f}s)", flush=True)
            outputs.append(
                {
                    "label": salida_word.name,
                    "url": f"/download?run={run_id}&file={salida_word.name}",
                }
            )

        print(f"[Run] Automatizacion completa. ({time.perf_counter() - started:.1f}s)", flush=True)
        return {"needs_confirmation": False, "files": outputs, "month": mes}
    finally:
        shutil.rmtree(run_upload_dir, ignore_errors=True)


class AppHandler(BaseHTTPRequestHandler):
    def log_message(self, format: str, *args):
        return

    def send_bytes(self, content: bytes, content_type: str, status: int = 200, extra_headers: dict[str, str] | None = None):
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(content)))
        if extra_headers:
            for key, value in extra_headers.items():
                self.send_header(key, value)
        self.end_headers()
        self.wfile.write(content)

    def send_json(self, payload: dict, status: int = 200):
        self.send_bytes(json.dumps(payload, ensure_ascii=False).encode("utf-8"), "application/json; charset=utf-8", status)

    def serve_file(self, path: Path):
        if not path.exists() or not path.is_file():
            self.send_bytes(b"No encontrado.", "text/plain; charset=utf-8", 404)
            return
        content_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        if path.suffix == ".css":
            content_type = "text/css; charset=utf-8"
        elif path.suffix == ".js":
            content_type = "application/javascript; charset=utf-8"
        elif path.suffix == ".html":
            content_type = "text/html; charset=utf-8"
        self.send_bytes(path.read_bytes(), content_type)

    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path == "/":
            self.serve_file(TEMPLATE_DIR / "index.html")
            return
        if parsed.path == "/xm":
            self.serve_file(TEMPLATE_DIR / "xm.html")
            return
        if parsed.path == "/helpdesk":
            self.serve_file(TEMPLATE_DIR / "helpdesk.html")
            return
        if parsed.path.startswith("/static/"):
            relative = safe_filename(Path(parsed.path).name, "")
            self.serve_file(STATIC_DIR / relative)
            return
        if parsed.path == "/download":
            params = parse_qs(parsed.query)
            run_id = safe_filename(params.get("run", [""])[0], "")
            filename = safe_filename(params.get("file", [""])[0], "")
            target = OUTPUT_DIR / run_id / filename
            if not run_id or not filename or not target.exists() or target.parent != OUTPUT_DIR / run_id:
                self.send_bytes(b"Archivo no encontrado.", "text/plain; charset=utf-8", 404)
                return
            mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"
            self.send_bytes(
                target.read_bytes(),
                mime,
                200,
                {"Content-Disposition": f'attachment; filename="{html.escape(filename)}"'},
            )
            try:
                target.unlink()
                if not any(target.parent.iterdir()):
                    target.parent.rmdir()
            except OSError:
                pass
            return
        self.send_bytes(b"No encontrado.", "text/plain; charset=utf-8", 404)

    def do_POST(self):
        path = urlparse(self.path).path
        if path not in {"/run", "/helpdesk"}:
            self.send_json({"ok": False, "error": "Ruta no encontrada."}, 404)
            return
        try:
            print(f"[HTTP] POST {path} recibido ({self.headers.get('Content-Length', '0')} bytes).", flush=True)
            step_started = time.perf_counter()
            fields, files = parse_multipart(self)
            print(f"[HTTP] Formulario parseado. ({time.perf_counter() - step_started:.1f}s)", flush=True)
            if path == "/helpdesk":
                result = run_helpdesk_automation(fields, files)
                print(f"[HTTP] Helpdesk terminado. Enviando respuesta con {len(result.get('files', []))} archivo(s).", flush=True)
                self.send_json({"ok": True, **result})
            else:
                result = run_automation_v2(fields, files)
                print(f"[HTTP] Automatizacion terminada. Enviando respuesta con {len(result.get('files', []))} archivo(s).", flush=True)
                self.send_json({"ok": True, **result})
            print("[HTTP] Respuesta enviada al navegador.", flush=True)
        except Exception as exc:
            traceback.print_exc()
            self.send_json({"ok": False, "error": str(exc)}, 400)


def main():
    shutil.rmtree(UPLOAD_DIR, ignore_errors=True)
    shutil.rmtree(OUTPUT_DIR, ignore_errors=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    host = "127.0.0.1"
    port = 8765
    server = ThreadingHTTPServer((host, port), AppHandler)
    print(f"Frontend disponible en http://{host}:{port}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nServidor detenido.")
    finally:
        server.server_close()


if __name__ == "__main__":
    main()

